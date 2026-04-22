import json
import re
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import RAGConfig
from app.services.llm_service import get_embeddings


class IndexingService:
    def __init__(self, chroma_client):
        self.chroma_client = chroma_client
        self.embeddings = get_embeddings()

    async def index_file(
        self,
        file_path: str,
        file_id: str,
        filename: str,
        file_extension: str,
        collection_name: str,
        user_id: str,
        splitting_method: str = "recursive",
    ) -> int:
        """Parse, chunk, embed, and store a file. Returns chunk count."""
        text = self._parse_file(file_path, file_extension)
        if not text.strip():
            raise ValueError("No text content could be extracted from the file")

        chunks = self._chunk_text(text, file_name=filename, method=splitting_method)
        if not chunks:
            raise ValueError("File produced no chunks after splitting")

        # Gap 1: Per-user collection
        user_collection = "system_knowledge"
        collection = self.chroma_client.get_or_create_collection(
            name=user_collection,
            metadata={"hnsw:space": "cosine"},
        )

        # Embed all chunks
        texts = [c.page_content for c in chunks]
        embeddings = self.embeddings.embed_documents(texts)

        # Gap 4: Rich metadata with section headings tracked as state
        metadatas = []
        current_heading = ""
        for i, chunk in enumerate(chunks):
            # Track heading across chunks
            for line in chunk.page_content.split("\n"):
                if line.startswith("## ") or line.startswith("# "):
                    current_heading = line.lstrip("#").strip()
                    break

            meta = {
                "file_id": file_id,
                "file_name": filename,
                "file_extension": file_extension,
                "chunk_index": i,
                "user_id": user_id,
                "section_heading": current_heading,
                "char_count": len(chunk.page_content),
                "doc_type": self._infer_doc_type(filename, file_extension),
                "is_whole_doc": chunk.metadata.get("is_whole_doc", False),
            }
            if "page" in chunk.metadata:
                meta["page_number"] = chunk.metadata["page"]
            metadatas.append(meta)

        ids = [f"{file_id}_chunk_{i}" for i in range(len(chunks))]

        collection.add(
            documents=texts,
            embeddings=embeddings,
            metadatas=metadatas,
            ids=ids,
        )

        self._refresh_kb_descriptor()

        return len(chunks)

    def delete_file_vectors(self, collection_name: str = "system_knowledge", file_id: str = ""):
        """Delete chunks for a specific file from the collection."""
        try:
            collection = self.chroma_client.get_collection(collection_name)
            collection.delete(where={"file_id": file_id})
        except (ValueError, Exception):
            pass
        self._refresh_kb_descriptor()

    def _refresh_kb_descriptor(self):
        """Regenerate the KB descriptor after mutation. Swallows errors to avoid breaking ingestion."""
        try:
            from app.services.rag_service import RAGService
            RAGService(self.chroma_client).rebuild_kb_descriptor()
        except Exception:
            import logging
            logging.getLogger(__name__).warning(
                "KB descriptor refresh failed; tool description may be stale.",
                exc_info=True,
            )

    def _parse_file(self, file_path: str, extension: str) -> str:
        """Parse a file into raw text."""
        ext = extension.lower().lstrip(".")
        path = Path(file_path)

        if ext in ("txt", "md"):
            return path.read_text(encoding="utf-8")
        if ext == "json":
            return self._parse_json(path)
        if ext == "pdf":
            return self._parse_pdf(path)
        if ext == "docx":
            return self._parse_docx(path)
        if ext in ("csv",):
            return self._parse_csv(path)
        if ext in ("xlsx", "xls"):
            return self._parse_excel(path)

        raise ValueError(f"Unsupported file type: .{ext}")

    def _parse_json(self, path: Path) -> str:
        data = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(data, str):
            return data
        return json.dumps(data, indent=2, ensure_ascii=False)

    def _parse_pdf(self, path: Path) -> str:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def _parse_docx(self, path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _parse_csv(self, path: Path) -> str:
        import pandas as pd
        df = pd.read_csv(str(path))
        return df.to_string(index=False)

    def _parse_excel(self, path: Path) -> str:
        import pandas as pd
        sheets = pd.read_excel(str(path), sheet_name=None)
        parts = []
        for sheet_name, df in sheets.items():
            parts.append(f"--- Sheet: {sheet_name} ---\n{df.to_string(index=False)}")
        return "\n\n".join(parts)

    def _chunk_text(self, text: str, file_name: str = "", method: str = "recursive") -> list:
        """Split text into chunks. Small files stored as single chunk (Gap 9).

        Headers are applied inside _chunk_recursive() — not here (avoids double headers).
        """
        # Gap 9: Skip chunking for small files WITHOUT structure.
        # If the file has ## headings, always section-split for better retrieval.
        word_count = len(text.split())
        has_headings = "\n## " in text or text.strip().startswith("## ")
        if word_count <= RAGConfig.SMALL_FILE_THRESHOLD and not has_headings:
            header = f"[Source: {file_name} | Whole document]\n"
            return [Document(
                page_content=header + text,
                metadata={"chunk_index": 0, "is_whole_doc": True},
            )]

        if method == "semantic":
            return self._chunk_semantic(text, file_name)
        return self._chunk_recursive(text, file_name)

    def _chunk_recursive(self, text: str, file_name: str = "") -> list:
        """Section-first chunking:
        1. Split by ## headings into sections
        2. Keep sections whole if they fit within CHUNK_SIZE
        3. Only sub-split oversized sections using paragraph/sentence fallbacks
        """
        # Prepend \n so a document starting with ## on line 1 is caught
        sections = re.split(r'(?=\n## )', '\n' + text.lstrip())

        chunks = []
        for section in sections:
            section = section.strip()
            if not section:
                continue

            if len(section) <= RAGConfig.CHUNK_SIZE:
                # Section fits — keep it whole
                chunks.append(Document(page_content=section, metadata={}))
            else:
                # Section too large — sub-split with heading-aware fallbacks
                sub_splitter = RecursiveCharacterTextSplitter(
                    separators=["\n### ", "\n#### ", "\n---", "\n\n", "\n", ". ", " "],
                    chunk_size=RAGConfig.CHUNK_SIZE,
                    chunk_overlap=RAGConfig.CHUNK_OVERLAP,
                )
                sub_chunks = sub_splitter.create_documents([section])
                chunks.extend(sub_chunks)

        # Fallback for plain text with no headings
        if not chunks and text.strip():
            fallback = RecursiveCharacterTextSplitter(
                separators=["\n\n", "\n", ". ", " "],
                chunk_size=RAGConfig.CHUNK_SIZE,
                chunk_overlap=RAGConfig.CHUNK_OVERLAP,
            )
            chunks = fallback.create_documents([text])

        # Gap 2: Contextual headers — applied here (single location)
        current_heading = ""
        for chunk in chunks:
            for line in chunk.page_content.split("\n"):
                if line.startswith("## ") or line.startswith("# "):
                    current_heading = line.lstrip("#").strip()
                    break

            header = f"[Source: {file_name}"
            if current_heading:
                header += f" | Section: {current_heading}"
            header += "]\n"
            chunk.page_content = header + chunk.page_content

        return chunks

    def _chunk_semantic(self, text: str, file_name: str = "") -> list:
        try:
            from langchain_experimental.text_splitter import SemanticChunker
            chunker = SemanticChunker(self.embeddings)
            chunks = chunker.create_documents([text])
            # Add contextual headers
            for chunk in chunks:
                chunk.page_content = f"[Source: {file_name}]\n" + chunk.page_content
            return chunks
        except ImportError:
            return self._chunk_recursive(text, file_name)

    @staticmethod
    def _infer_doc_type(file_name: str, ext: str) -> str:
        ext = ext.lower().lstrip(".")
        if ext in ("pdf",):
            return "document"
        if ext in ("csv", "xlsx", "xls"):
            return "data"
        if ext in ("md", "txt"):
            return "text"
        if ext in ("json",):
            return "structured"
        if ext in ("docx",):
            return "document"
        return "other"
